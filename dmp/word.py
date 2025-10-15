from docx import Document

doc = Document()

doc.add_heading('Data Management Plan', level=1)

doc.add_paragraph('Action Number: [10120595]')
doc.add_paragraph('Action Acronym: [HIP-HP]')
doc.add_paragraph('Action Title: [Probing nuclear parton dynamics in Heavy-Ion Collisions with Hard Probes]')
doc.add_paragraph('Data: [23/09/2025]')
doc.add_paragraph('DMP version: [1.0]')

doc.add_heading('1. Data Summary', level=2)

doc.add_paragraph('')

output_path = './dmp.docx'
doc.save(output_path)